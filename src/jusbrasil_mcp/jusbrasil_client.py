"""Cliente JusBrasil via Scrapling (StealthyFetcher) — autenticado, headless.

Migrado de patchright para o StealthyFetcher do Scrapling, que passa o Cloudflare
em HEADLESS puro. Mantém o login com credenciais (Keychain) para acessar o
INTEIRO TEOR das decisões — capacidade central do projeto.

Tools cobertas:
  - buscar_jurisprudencia / ler_decisao  (metadados + ementa)
  - ler_inteiro_teor                     (texto integral do acórdão; exige login)

O login é automático e idempotente: se a sessão persistente já estiver válida,
pula o login; senão, preenche o formulário (2 etapas) com as credenciais do
Keychain. Sem credenciais, opera anônimo (números podem vir mascarados; o
inteiro teor fica indisponível).
"""
from __future__ import annotations

import asyncio
import json
import re
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Optional
from urllib.parse import quote_plus

import keyring
from bs4 import BeautifulSoup
from scrapling.fetchers import StealthyFetcher

SERVICO = "mcp-jusbrasil"
URL_BASE = "https://www.jusbrasil.com.br"
URL_LOGIN = f"{URL_BASE}/login"
URL_BUSCA = f"{URL_BASE}/jurisprudencia/busca"
USER_DATA_DIR = Path.home() / ".mcp-jusbrasil-scrapling-profile"


@dataclass
class Resultado:
    titulo: str
    tribunal: Optional[str]
    numero_cnj: Optional[str]
    data: Optional[str]
    ementa: Optional[str]
    url: Optional[str]

    def to_dict(self) -> dict:
        return asdict(self)



# Mapa de tribunais: usuario pode dizer "TJPI", "TJ-PI", "tj pi",
# "Tribunal de Justiça do Piauí" -> convertemos pro slug "tj_pi"
@dataclass
class Decisao:
    """Metadados completos de uma decisao individual."""
    url: str
    titulo: Optional[str]
    tribunal: Optional[str]
    tipo: Optional[str]
    numero_cnj: Optional[str]
    relator: Optional[str]
    data_julgamento: Optional[str]
    data_publicacao: Optional[str]
    orgao_julgador: Optional[str]
    ementa: Optional[str]

    def to_dict(self) -> dict:
        return asdict(self)

    def citacao_abnt(self) -> str:
        """Monta citacao por tribunal. STJ usa formato proprio (recurso + registro)."""
        if (self.tribunal or "").upper() == "STJ":
            return self._citacao_stj()
        return self._citacao_default()

    def _citacao_default(self) -> str:
        """Formato generico para TJs, TRFs, TST."""
        head = ""
        if self.tribunal:
            head = self.tribunal
        if self.tipo:
            head = f"{head} - {self.tipo}" if head else self.tipo
        if self.numero_cnj:
            head = f"{head}: {self.numero_cnj}" if head else self.numero_cnj
        extras = []
        if self.relator:
            extras.append(f"Relator: {self.relator}")
        if self.data_julgamento:
            extras.append(f"Data de Julgamento: {self.data_julgamento}")
        if self.orgao_julgador:
            extras.append(self.orgao_julgador)
        joined = ", ".join([p for p in [head, *extras] if p])
        return f"({joined})" if joined else ""

    def _citacao_stj(self) -> str:
        """Formato STJ convencional: (STJ - AgInt no REsp 1.846.222/RS (2019/0326486-1),
        Rel. Min. Luis Felipe Salomao, Quarta Turma, julgado em DD/MM/YYYY, DJe DD/MM/YYYY)."""
        recurso = None
        if self.titulo:
            m = re.search(r":\s*(.+?)\s+(\d+)\s+([A-Z]{2})\s+(\d{4}/\d{7}-\d)\s*$", self.titulo)
            if m:
                tipo_rec, num, uf, registro = m.group(1).strip(), m.group(2), m.group(3), m.group(4)
                num_fmt = f"{int(num):,}".replace(",", ".")
                recurso = f"{tipo_rec} {num_fmt}/{uf} ({registro})"

        if recurso:
            head = f"STJ - {recurso}"
        elif self.numero_cnj:
            head = f"STJ - {self.numero_cnj}"
        else:
            head = "STJ"

        extras = []
        if self.relator:
            extras.append(f"Rel. Min. {_title_case_pt(self.relator)}")
        if self.orgao_julgador:
            extras.append(_orgao_stj_extenso(self.orgao_julgador))
        if self.data_julgamento:
            extras.append(f"julgado em {self.data_julgamento}")
        if self.data_publicacao:
            dp = self.data_publicacao
            if not dp.lower().startswith("dje"):
                dp = f"DJe {dp}"
            extras.append(dp)

        return f"({', '.join([head] + extras)})"


@dataclass
class InteiroTeor:
    """Texto integral de uma decisao (relatorio + voto + acordao)."""
    url: str
    url_inteiro_teor: Optional[str]
    texto: Optional[str]
    n_caracteres: int
    autenticado: bool

    def to_dict(self) -> dict:
        return asdict(self)


def formatar_cnj(numero_raw):
    """Converte numero cru em formato CNJ.
    Ex.: '7034194820198180000' -> '0703419-48.2019.8.18.0000'
    """
    if not numero_raw:
        return None
    digitos = re.sub(r"\D", "", str(numero_raw))
    if not digitos:
        return None
    digitos = digitos.zfill(20)
    if len(digitos) != 20:
        return None
    m = re.match(r"(\d{7})(\d{2})(\d{4})(\d{1})(\d{2})(\d{4})", digitos)
    if not m:
        return None
    return f"{m.group(1)}-{m.group(2)}.{m.group(3)}.{m.group(4)}.{m.group(5)}.{m.group(6)}"


_SIGLAS_ESPECIAIS = {
    "STF", "STJ", "TST", "TSE", "STM", "TCU", "TNU", "TRU", "CNJ", "CARF"
}

_NOMES_EXTENSOS = {
    "tribunal de justica do piaui": "tj_pi",
    "tribunal de justica de sao paulo": "tj_sp",
    "tribunal de justica do rio de janeiro": "tj_rj",
    "tribunal de justica de minas gerais": "tj_mg",
    "tribunal de justica do rio grande do sul": "tj_rs",
    "supremo tribunal federal": "stf",
    "superior tribunal de justica": "stj",
    "tribunal superior do trabalho": "tst",
    "tribunal superior eleitoral": "tse",
}


def _normalizar_tribunal(tribunal: str) -> str:
    """Converte qualquer jeito que o usuario escreva num slug do JusBrasil."""
    import unicodedata
    if not tribunal:
        return ""
    raw = unicodedata.normalize("NFKD", tribunal).encode("ascii", "ignore").decode().lower().strip()
    if raw in _NOMES_EXTENSOS:
        return _NOMES_EXTENSOS[raw]
    t = re.sub(r"[\s\-.]+", "", tribunal.upper())
    if t in _SIGLAS_ESPECIAIS:
        return t.lower()
    m = re.match(r"^(TJ|TRF|TRT|TRE|TJM|TCE)([A-Z0-9]+)$", t)
    if m:
        return f"{m.group(1).lower()}_{m.group(2).lower()}"
    return re.sub(r"[\s\-.]+", "_", tribunal.lower())


def _strip_acentos(s: str) -> str:
    import unicodedata
    return unicodedata.normalize("NFKD", s).encode("ascii", "ignore").decode()


def _norm_tribunal_filtro(tribunal: str) -> str:
    """Valor do filtro `tribunal` (sigla minuscula). Aceita virgula p/ varios.
    Ex.: "STJ" -> "stj"; "STJ, STF" -> "stj,stf"; "TJs" -> "tjs".
    """
    raw = _strip_acentos(tribunal).lower().strip()
    if "," in raw:
        return ",".join(_norm_tribunal_filtro(t) for t in raw.split(",") if t.strip())
    return re.sub(r"\s+", "", raw)


_JURIS_TYPES = {
    "acordao": "acordao", "acordaos": "acordao",
    "sumula": "sumula", "sumulas": "sumula",
    "decisao": "decisao", "decisaomonocratica": "decisao", "monocratica": "decisao",
    "sentenca": "sentenca", "sentencas": "sentenca",
    "despacho": "despacho", "despachos": "despacho",
    "orientacao": "orientacao", "orientacaojurisprudencial": "orientacao",
}


def _map_juris_type(tipo: str) -> Optional[str]:
    k = re.sub(r"\W", "", _strip_acentos(tipo).lower())
    return _JURIS_TYPES.get(k)


def _map_periodo(periodo: str) -> Optional[str]:
    """Valor do filtro de periodo `l` (ex.: "365dias"). Aceita apelidos e n de dias."""
    p = re.sub(r"\s+", "", _strip_acentos(str(periodo)).lower())
    apelidos = {
        "ultimasemana": "7dias", "semana": "7dias",
        "ultimomes": "30dias", "mes": "30dias",
        "ultimoano": "365dias", "ano": "365dias", "1ano": "365dias", "ultimos12meses": "365dias",
    }
    if p in apelidos:
        return apelidos[p]
    if re.fullmatch(r"\d+dias", p):
        return p
    m = re.fullmatch(r"(\d+)(?:d|dia|dias)?", p)
    return f"{m.group(1)}dias" if m else None


class JusBrasilClient:
    """Fetch via Scrapling StealthyFetcher; login automatico por credenciais.

    Mantem a interface async (start/close/buscar/ler) do server.py. start()/close()
    sao no-op (o StealthyFetcher gerencia o proprio navegador por chamada).
    """

    def __init__(self, headless: bool = True, timeout_ms: int = 60_000):
        self.headless = headless
        self.timeout_ms = max(timeout_ms, 60_000)
        self._logged_in = False
        self._login_lock = asyncio.Lock()
        self.email = keyring.get_password(SERVICO, "email")
        self.senha = keyring.get_password(SERVICO, "senha")
        USER_DATA_DIR.mkdir(exist_ok=True)

    async def start(self) -> None:
        return  # Scrapling lanca o navegador por fetch

    async def close(self) -> None:
        return

    def _fetch_kwargs(self, page_action=None) -> dict:
        kw = dict(
            headless=self.headless,
            real_chrome=True,
            solve_cloudflare=True,
            user_data_dir=str(USER_DATA_DIR),
            network_idle=True,
            google_search=True,
            timeout=self.timeout_ms,
        )
        if page_action is not None:
            kw["page_action"] = page_action
        return kw

    # ---------- Login ----------

    def _login_action(self, page):
        """Roda DENTRO do navegador (page_action sincrono). Idempotente:
        se a sessao persistente ja redirecionou pra fora de /login, nao faz nada."""
        page.wait_for_timeout(2500)
        if "/login" not in (page.url or "").lower():
            print("[LOGIN] sessao persistente valida — pulando login")
            return
        for sel in ['button:has-text("Aceitar cookies")', 'button:has-text("Aceitar todos")',
                    'button:has-text("Concordar")']:
            try:
                b = page.locator(sel).first
                if b.is_visible(timeout=600):
                    b.click(timeout=1200)
                    break
            except Exception:
                pass
        try:
            print("[LOGIN 1/2] email...")
            page.locator('input[type="email"]').first.fill(self.email, timeout=8000)
            page.locator('button[type="submit"]').first.click(timeout=4000)
            page.wait_for_timeout(4500)
            print("[LOGIN 2/2] senha...")
            page.locator('input[type="password"]').first.fill(self.senha, timeout=12000)
            page.locator('button[type="submit"]').first.click(timeout=4000)
            try:
                page.wait_for_load_state("networkidle", timeout=25000)
            except Exception:
                pass
            page.wait_for_timeout(3000)
            print(f"[LOGIN] apos submit -> {page.url}")
        except Exception as e:
            print(f"[LOGIN] aviso durante login: {e}")

    async def _ensure_login(self) -> bool:
        """Garante sessao autenticada (uma vez por processo). Sem credenciais: anonimo."""
        if self._logged_in:
            return True
        if not (self.email and self.senha):
            return False
        async with self._login_lock:
            if self._logged_in:
                return True
            print("[LOGIN] autenticando no JusBrasil...")
            await asyncio.to_thread(
                StealthyFetcher.fetch, URL_LOGIN, **self._fetch_kwargs(page_action=self._login_action)
            )
            self._logged_in = True
            return True

    async def _fetch_html(self, url: str, autenticar: bool = True) -> str:
        if autenticar:
            await self._ensure_login()
        page = await asyncio.to_thread(StealthyFetcher.fetch, url, **self._fetch_kwargs())
        return page.html_content

    # ---------- Busca ----------

    async def buscar_jurisprudencia(
        self,
        query: str,
        limite: int = 10,
        tribunal: Optional[str] = None,
        tipo: Optional[str] = None,
        periodo: Optional[str] = None,
        ordenacao: Optional[str] = None,
    ) -> list[Resultado]:
        if not query or not query.strip():
            raise ValueError("query vazia")
        url = self._montar_url_busca(query, tribunal, tipo, periodo, ordenacao)
        print(f"[BUSCA] {url}")
        html = await self._fetch_html(url)
        return self._parse_resultados(html, limite)

    @staticmethod
    def _montar_url_busca(query, tribunal=None, tipo=None, periodo=None, ordenacao=None,
                          juris_type_forcado=None) -> str:
        """Monta a URL de busca com os filtros REAIS do JusBrasil (descobertos via UI).
        tribunal=<sigla minuscula> (stf/stj/tst/tjs/...); jurisType=<tipo>; l=<N>dias; o=data.
        """
        params = [f"q={quote_plus(query)}"]
        if tribunal:
            params.append(f"tribunal={quote_plus(_norm_tribunal_filtro(tribunal))}")
        jt = juris_type_forcado or (_map_juris_type(tipo) if tipo else None)
        if jt:
            params.append(f"jurisType={jt}")
        if periodo:
            l = _map_periodo(periodo)
            if l:
                params.append(f"l={l}")
        if ordenacao and re.sub(r"\W", "", ordenacao.lower()) in ("data", "recente", "recentes", "maisrecentes"):
            params.append("o=data")
        return f"{URL_BUSCA}?{'&'.join(params)}"

    def _parse_resultados(self, html: str, limite: int) -> list[Resultado]:
        soup = BeautifulSoup(html, "html.parser")
        candidatos = soup.select("article, li[data-testid*='result'], div.SearchResult")
        if not candidatos:
            candidatos = [a.parent for a in soup.select('a[href*="/jurisprudencia/"]') if a.parent]

        vistos: set[str] = set()
        resultados: list[Resultado] = []
        for node in candidatos:
            if len(resultados) >= limite:
                break
            try:
                link = node.find("a", href=True)
                href = link["href"] if link else None
                if href and href.startswith("/"):
                    href = f"{URL_BASE}{href}"
                if not href or "/jurisprudencia/" not in href or href in vistos:
                    continue
                vistos.add(href)
                titulo = self._first_text(node, ["h2", "h3", "h4", "a"]) or "(sem titulo)"
                import re as _re
                m = _re.match(r"^([A-Z][A-Z0-9-]{1,8}(?:\s*[A-Z0-9-]+)?)\s*[-–]", titulo)
                tribunal = m.group(1).strip() if m else self._first_text(node, [".tribunal", "span.source"])
                num_raw = None
                m_num = re.search(r"(\d{15,20})", titulo)
                if m_num:
                    num_raw = m_num.group(1)
                numero_cnj = formatar_cnj(num_raw) if num_raw else None
                resultados.append(
                    Resultado(
                        titulo=titulo,
                        tribunal=tribunal,
                        numero_cnj=numero_cnj,
                        data=self._first_text(node, [".data", "[data-testid*='date']", "time"]),
                        ementa=self._first_text(node, [".ementa", "[data-testid*='ementa']", "p"]),
                        url=href,
                    )
                )
            except Exception as e:
                print(f"[PARSE] erro num item: {e}")
                continue
        return resultados

    # ---------- Sumulas ----------

    async def buscar_sumulas(
        self, query: str, limite: int = 10, tribunal: Optional[str] = None
    ) -> list[Resultado]:
        """Busca sumulas no JusBrasil usando o filtro real jurisType=sumula."""
        if not query or not query.strip():
            raise ValueError("query vazia")
        url = self._montar_url_busca(query, tribunal=tribunal, juris_type_forcado="sumula")
        print(f"[SUMULA] {url}")
        html = await self._fetch_html(url)
        return self._parse_resultados(html, limite)

    # ---------- Decisao ----------

    async def ler_decisao(self, url: str):
        """Abre pagina de decisao e extrai todos os metadados."""
        if not url or "jusbrasil.com.br" not in url:
            raise ValueError(f"URL invalida: {url}")
        print(f"[LER] {url}")
        html = await self._fetch_html(url)
        dec = self._parse_decisao(html, url)
        if not dec.numero_cnj:
            real = self._cnj_real_do_apollo(html)
            if real:
                dec.numero_cnj = real
        return dec

    # ---------- Dossie ----------

    async def compilar_dossie(
        self,
        urls: list[str],
        incluir_inteiro_teor: bool = False,
        titulo: Optional[str] = None,
        caminho: Optional[str] = None,
    ) -> dict:
        """Compila varias decisoes/sumulas num unico documento .docx."""
        if not urls:
            raise ValueError("lista de urls vazia")
        titulo = titulo or "Dossiê de Jurisprudência"
        itens = []
        for u in urls:
            try:
                d = await self.ler_decisao(u)
            except Exception as e:
                itens.append({"url": u, "erro": str(e)[:150]})
                continue
            teor = None
            if incluir_inteiro_teor:
                try:
                    teor = (await self.ler_inteiro_teor(u)).texto
                except Exception:
                    teor = None
            itens.append({"url": u, "decisao": d, "inteiro_teor": teor})
        arquivo = self._gerar_docx(titulo, itens, caminho)
        return {
            "arquivo": arquivo,
            "n_itens": sum(1 for i in itens if "decisao" in i),
            "n_falhas": sum(1 for i in itens if "erro" in i),
            "com_inteiro_teor": bool(incluir_inteiro_teor),
        }

    @staticmethod
    def _gerar_docx(titulo: str, itens: list[dict], caminho: Optional[str]) -> str:
        from datetime import datetime
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

        doc.add_heading(titulo, level=0)
        meta = doc.add_paragraph(
            f"Gerado em {datetime.now():%d/%m/%Y %H:%M} · {len(itens)} item(ns) · "
            "Fonte: JusBrasil (agregador privado; confira no site oficial do tribunal)."
        )
        meta.runs[0].italic = True

        for i, item in enumerate(itens, 1):
            if "erro" in item:
                doc.add_heading(f"{i}. (falha ao ler)", level=1)
                doc.add_paragraph(item["url"])
                doc.add_paragraph(f"Erro: {item['erro']}")
                continue
            d = item["decisao"]
            cab = " ".join(x for x in [d.tribunal, d.tipo] if x) or "Decisão"
            doc.add_heading(f"{i}. {cab}", level=1)
            cit = d.citacao_abnt()
            if cit:
                p = doc.add_paragraph()
                p.add_run(cit).italic = True
            if d.ementa:
                doc.add_heading("Ementa", level=2)
                pe = doc.add_paragraph(d.ementa)
                pe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if item.get("inteiro_teor"):
                doc.add_heading("Inteiro teor", level=2)
                for par in re.split(r"\n{2,}", item["inteiro_teor"]):
                    par = par.strip()
                    if par:
                        pp = doc.add_paragraph(par)
                        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            link = d.url or item["url"]
            pl = doc.add_paragraph()
            pl.add_run("Verificar autenticidade: ").bold = True
            pl.add_run(link)

        if not caminho:
            import unicodedata
            base = unicodedata.normalize("NFKD", titulo).encode("ascii", "ignore").decode().lower()
            slug = re.sub(r"[^a-z0-9]+", "-", base).strip("-")[:40] or "dossie"
            caminho = str(Path.home() / "Downloads" / f"{slug}-{datetime.now():%Y%m%d-%H%M}.docx")
        Path(caminho).expanduser().parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(Path(caminho).expanduser()))
        return str(Path(caminho).expanduser())

    # ---------- Inteiro teor ----------

    async def ler_inteiro_teor(self, url: str) -> InteiroTeor:
        """Extrai o TEXTO INTEGRAL de uma decisao (relatorio+voto+acordao).

        Aceita a URL da decisao OU a URL direta de inteiro teor. Exige login
        (credenciais no Keychain) — sem sessao autenticada o conteudo nao vem.
        """
        if not url or "jusbrasil.com.br" not in url:
            raise ValueError(f"URL invalida: {url}")
        autenticado = await self._ensure_login()

        url_teor = url
        if "/inteiro-teor-" not in url:
            print(f"[TEOR] localizando inteiro teor a partir da decisao: {url}")
            html_dec = await self._fetch_html(url)
            url_teor = self._achar_url_inteiro_teor(html_dec, url)
            if not url_teor:
                return InteiroTeor(url=url, url_inteiro_teor=None, texto=None,
                                   n_caracteres=0, autenticado=autenticado)

        print(f"[TEOR] {url_teor}")
        html_teor = await self._fetch_html(url_teor)
        texto = self._extrair_inteiro_teor(html_teor)
        return InteiroTeor(
            url=url, url_inteiro_teor=url_teor, texto=texto,
            n_caracteres=len(texto or ""), autenticado=autenticado,
        )

    @staticmethod
    def _achar_url_inteiro_teor(html_decisao: str, url_decisao: str) -> Optional[str]:
        m = re.search(r'https://www\.jusbrasil\.com\.br/jurisprudencia/[a-z0-9_-]+/\d+/inteiro-teor-\d+',
                      html_decisao)
        if m:
            return m.group(0)
        m = re.search(r'(/jurisprudencia/[a-z0-9_-]+/\d+/inteiro-teor-\d+)', html_decisao)
        if m:
            return f"{URL_BASE}{m.group(1)}"
        m = re.search(r'/inteiro-teor-\d+', html_decisao)
        if m:
            return url_decisao.rstrip("/") + m.group(0)
        return None

    @staticmethod
    def _extrair_inteiro_teor(html: str) -> Optional[str]:
        """Texto integral fica no campo decisionHtml do no INTEIRO_TEOR do Apollo."""
        m = re.search(r'<script id="__NEXT_DATA__"[^>]*>(.+?)</script>', html, re.S)
        if not m:
            return None
        try:
            apollo = json.loads(m.group(1)).get("props", {}).get("pageProps", {}).get("__APOLLO_STATE__", {})
        except json.JSONDecodeError:
            return None
        for k, v in apollo.items():
            if k.startswith("Document:") and isinstance(v, dict) and v.get("decisionHtml"):
                texto = BeautifulSoup(v["decisionHtml"], "html.parser").get_text("\n", strip=True)
                texto = re.sub(r"\n{3,}", "\n\n", texto).strip()
                return texto or None
        return None

    def _parse_next_data(self, html: str) -> Optional[dict]:
        m = re.search(r'<script id="__NEXT_DATA__"[^>]*>(.+?)</script>', html, re.S)
        if not m:
            return None
        try:
            data = json.loads(m.group(1))
        except json.JSONDecodeError:
            return None
        apollo = data.get("props", {}).get("pageProps", {}).get("__APOLLO_STATE__", {})
        for key, val in apollo.items():
            if key.startswith("Document:") and isinstance(val, dict) and "decisionDepartament" in val:
                return val
        return None

    def _cnj_real_do_apollo(self, html: str) -> Optional[str]:
        """CNJ real (informacoes_gerais.numero_processo) p/ quando decisionLabel vem mascarado."""
        m = re.search(r'<script id="__NEXT_DATA__"[^>]*>(.+?)</script>', html, re.S)
        if not m:
            return None
        try:
            apollo = json.loads(m.group(1)).get("props", {}).get("pageProps", {}).get("__APOLLO_STATE__", {})
        except json.JSONDecodeError:
            return None
        for key, val in apollo.items():
            if not (key.startswith("Document:") and isinstance(val, dict)):
                continue
            for k, v in val.items():
                if "enrichedContents" in k and isinstance(v, list):
                    for item in v:
                        payload = (item or {}).get("payload", {}) if isinstance(item, dict) else {}
                        geral = payload.get("informacoes_gerais", {}) if isinstance(payload, dict) else {}
                        num = geral.get("numero_processo")
                        if num and "X" not in num.upper():
                            return num.strip()
        return None

    def _parse_decisao(self, html: str, url: str):
        """Parser tunado com base no HTML real do JusBrasil."""
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text(" ", strip=True)

        def extrair(pattern, texto=text, flags=re.IGNORECASE):
            m = re.search(pattern, texto, flags)
            if not m:
                return None
            v = m.group(1).strip()
            v = re.sub(r"[\s,;\.·]+$", "", v)
            return v or None

        titulo = None
        h1 = soup.find(["h1", "h2"])
        if h1:
            titulo = h1.get_text(" ", strip=True)
        if not titulo:
            og = soup.find("meta", property="og:title")
            if og and og.get("content"):
                titulo = og["content"].strip()

        tribunal = None
        m_trib = re.search(r"/jurisprudencia/([a-z0-9_-]+)/", url)
        if m_trib:
            slug = m_trib.group(1)
            tribunal = slug.upper().replace("_", "-")

        tipo = extrair(
            r"(Apela(?:c|ç)(?:a|ã)o C(?:i|í)vel|Apela(?:c|ç)(?:a|ã)o"
            r"|Recurso Especial|Recurso Extraordin(?:a|á)rio|Habeas Corpus"
            r"|Agravo(?: de Instrumento| Interno| Regimental)?"
            r"|Embargos (?:de Declara(?:c|ç)(?:a|ã)o|Infringentes)"
            r"|Mandado de Seguran(?:c|ç)a)",
            titulo or text
        )

        nd = self._parse_next_data(html)
        if nd:
            from datetime import datetime
            cnj = None
            m_cnj = re.search(r"(\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4})", nd.get("decisionLabel") or "")
            if m_cnj:
                cnj = m_cnj.group(1)
            data_jul = None
            ts = nd.get("judgementDate")
            if ts:
                data_jul = datetime.fromtimestamp(ts / 1000).strftime("%d/%m/%Y")
            data_pub = None
            pubs = nd.get("decisionPublications") or []
            if pubs:
                desc = pubs[0].get('description({"stripTags":true})') or pubs[0].get("description")
                if desc:
                    data_pub = desc.strip()
            ementa_nd = None
            df = nd.get("decisionFacts") or {}
            if df:
                header = re.sub(r"<[^>]+>", "", df.get("header", "") or "")
                paragraphs = df.get("paragraphs") or []
                joined = (header + "\n" + "\n".join(paragraphs)).strip()
                ementa_nd = joined or None
            relator_nd = nd.get("chairman")
            if relator_nd:
                relator_nd = re.sub(r"^Ministr[oa]\s+", "", relator_nd).strip()
            return Decisao(
                url=url,
                titulo=nd.get("title") or titulo,
                tribunal=tribunal,
                tipo=tipo,
                numero_cnj=cnj,
                relator=relator_nd,
                data_julgamento=data_jul,
                data_publicacao=data_pub,
                orgao_julgador=nd.get("decisionDepartament"),
                ementa=ementa_nd,
            )

        numero_cnj = extrair(r"(\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4})")
        if not numero_cnj:
            mraw = re.search(r"(\d{19,20})", text)
            if mraw:
                numero_cnj = formatar_cnj(mraw.group(1))

        relator = extrair(
            r"Relator(?:a)?[\s·:\.]+([A-Z][A-Za-zÀ-ſ\s\.]{3,80}?)"
            r"(?=\s+(?:Julgado|Data|Órg|Org|C(?:a|â)mara|Ementa)|$)"
        )

        data_julgamento = extrair(r"Julgado em (\d{1,2}/\d{1,2}/\d{2,4})")
        if not data_julgamento:
            data_julgamento = extrair(r"Data de Julgamento[:\s]+(\d{1,2}/\d{1,2}/\d{2,4})")

        data_publicacao = extrair(r"(?:Publicado em|Data de Publica(?:c|ç)(?:a|ã)o)[:\s]+(\d{1,2}/\d{1,2}/\d{2,4})")

        orgao = extrair(
            r"(\d+(?:ª|a)\s+(?:C(?:a|â)mara|Turma)[^·\n]*?)"
            r"(?=\s+Relator|\s+Julgado|$)"
        )
        if not orgao:
            orgao = extrair(
                r"(?:Órg(?:a|ã)o Julgador|Classe)[:\s]+([^,\n·]{3,120}?)"
                r"(?=\s+Relator|\s+Julgado|\s+Data|Ementa|$)"
            )

        ementa = None
        m_em = re.search(
            r"Ementa[:\s]+"
            r"((?:APELA|RECURSO|A(?:C|Ç)(?:A|Ã)O|HABEAS|AGRAVO|EMBARGOS|"
            r"MANDADO|RECLAMA|CONFLITO|PROCESSO|REMESSA|REEXAME)[^$]+?)"
            r"(?=\s+(?:Marca Jus IA|Documentos anexos|Inteiro Teor|"
            r"Jus IA Pergunt|Compartilhar|Salvar|Perguntas e respostas)|$)",
            text, re.S | re.I
        )
        if m_em:
            ementa = re.sub(r"\s+", " ", m_em.group(1)).strip()[:4000]

        return Decisao(
            url=url, titulo=titulo, tribunal=tribunal, tipo=tipo,
            numero_cnj=numero_cnj, relator=relator,
            data_julgamento=data_julgamento, data_publicacao=data_publicacao,
            orgao_julgador=orgao, ementa=ementa,
        )

    @staticmethod
    def _first_text(node, selectors: list[str]) -> Optional[str]:
        for sel in selectors:
            el = node.select_one(sel)
            if el and el.get_text(strip=True):
                return el.get_text(" ", strip=True)
        return None


def _title_case_pt(name: str) -> str:
    out = name.title()
    for word in [" Do ", " Da ", " De ", " Dos ", " Das ", " E "]:
        out = out.replace(word, word.lower())
    return out


def _orgao_stj_extenso(orgao: str) -> str:
    m = re.match(r"^(T[1-6]|S[1-3]|CE)\s*-\s*(.+)$", orgao or "", re.I)
    if m:
        return m.group(2).title()
    return orgao or ""
